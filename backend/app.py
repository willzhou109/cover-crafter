import io
import os
import time
import threading
import logging
from typing import Dict, List, Tuple

from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from langchain_openai import ChatOpenAI
from pydantic import BaseModel
from PyPDF2 import PdfReader
import docx
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

load_dotenv()

MAX_FILE_SIZE = 5 * 1024 * 1024
MAX_TEXT_CHARS = 20000
RATE_LIMIT_COUNT = 10
RATE_LIMIT_WINDOW_SECONDS = 60

app = FastAPI()

_frontend_origins_env = os.getenv("FRONTEND_ORIGINS", "")
if _frontend_origins_env.strip():
    FRONTEND_ORIGINS = [origin.strip() for origin in _frontend_origins_env.split(",") if origin.strip()]
else:
    FRONTEND_ORIGINS = ["http://localhost:5173"]

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("cover_letter_api")

_rate_limit_lock = threading.Lock()
_rate_limit_hits: Dict[str, List[float]] = {}

app.add_middleware(
    CORSMiddleware,
    allow_origins=FRONTEND_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class CoverLetterPayload(BaseModel):
    coverLetter: str


def _get_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def _read_upload(upload: UploadFile) -> Tuple[str, bytes]:
    if not upload.filename:
        raise HTTPException(status_code=400, detail="Missing resume file name")

    data = upload.file.read()
    if len(data) == 0:
        raise HTTPException(status_code=400, detail="Empty resume file")
    if len(data) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Resume file exceeds 5MB limit")

    return upload.filename, data


def _extract_text(filename: str, data: bytes) -> str:
    ext = _get_extension(filename)

    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        text = "".join(page.extract_text() or "" for page in reader.pages)
        return text

    if ext == ".docx":
        doc = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    if ext == ".txt":
        return data.decode("utf-8", errors="ignore")

    raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT")


def _truncate_text(text: str) -> str:
    return text[:MAX_TEXT_CHARS]


def _build_prompt(resume_text: str, job_text: str) -> str:
    return f"""

You are an expert career writer.

Your task is to write a concise, professional, and tailored cover letter using the resume and job description below.

Before writing:
1. Identify the 4–6 most important skills or qualifications from the job description.
2. Select the most relevant experiences from the resume that demonstrate those skills.
3. Focus on alignment and impact rather than listing responsibilities.

Writing requirements:
- Length: 3–4 paragraphs (one page maximum)
- Begin the letter with "Dear Hiring Manager"
- Tone: professional, confident, and clear
- Match keywords and priorities from the job description naturally
- Summarize resume content; do NOT restate it verbatim
- Do NOT repeat any numbers or metrics already present in the resume
- Avoid generic phrases (e.g., “hard-working,” “fast learner”)
- Do not include the date

Header formatting (must appear at the top in this exact order):
Applicant information (each on its own line):
Full Name  
City of Residence (if available)
Email Address (if available)
Phone Number (if available)

Then include a separator line consisting of three hyphens:
---

Company information (each on its own line):
Company Name  
Company Address (if available)

Letter structure:
- Paragraph 1: Express interest in the role and briefly explain why you are a strong fit
- Paragraph 2: Highlight the most relevant technical skills and experiences that match the role
- Paragraph 3 (optional 4th): Emphasize alignment with the company/team and conclude professionally

Closing format:
- End the letter with the word "Sincerely,"
- Then include the applicant’s full name on the next line

Output rules:
- Output only the final cover letter text
- No bullet points
- No explanations or analysis

--- Resume ---
{resume_text}

--- Job Description ---
{job_text}
"""


def _get_llm() -> ChatOpenAI:
    if not os.getenv("OPENAI_API_KEY"):
        raise HTTPException(status_code=500, detail="Missing OPENAI_API_KEY in environment")
    return ChatOpenAI(model="gpt-4o", temperature=0.7)


@app.middleware("http")
async def log_and_rate_limit(request, call_next):
    start = time.monotonic()
    client_ip = request.client.host if request.client else "unknown"
    response = None

    if request.url.path == "/api/generate":
        now = time.monotonic()
        with _rate_limit_lock:
            hits = _rate_limit_hits.setdefault(client_ip, [])
            hits[:] = [ts for ts in hits if now - ts < RATE_LIMIT_WINDOW_SECONDS]
            if len(hits) >= RATE_LIMIT_COUNT:
                logger.warning("rate_limit block ip=%s path=%s", client_ip, request.url.path)
                raise HTTPException(status_code=429, detail="Rate limit exceeded. Try again later.")
            hits.append(now)

    try:
        response = await call_next(request)
    except HTTPException:
        raise
    finally:
        duration_ms = (time.monotonic() - start) * 1000
        status_code = getattr(response, "status_code", 500)
        logger.info(
            "request ip=%s method=%s path=%s status=%s duration_ms=%.1f",
            client_ip,
            request.method,
            request.url.path,
            status_code,
            duration_ms,
        )
    return response


@app.post("/api/generate")
async def generate_cover_letter(
    resume: UploadFile = File(...),
    jobText: str = Form(...),
):
    if not jobText.strip():
        raise HTTPException(status_code=400, detail="Missing job description")

    filename, data = _read_upload(resume)
    extracted = _extract_text(filename, data)
    if not extracted.strip():
        raise HTTPException(status_code=400, detail="Failed to extract resume text")

    resume_text = _truncate_text(extracted)
    job_text = _truncate_text(jobText.strip())

    prompt = _build_prompt(resume_text, job_text)
    llm = _get_llm()

    try:
        response = llm.invoke(prompt)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"OpenAI error: {exc}") from exc

    return {"coverLetter": response.content}


@app.post("/api/export/docx")
async def export_docx(payload: CoverLetterPayload):
    text = payload.coverLetter.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Missing cover letter text")

    doc = docx.Document()
    for para in text.split("\n\n"):
        doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    headers = {"Content-Disposition": "attachment; filename=cover_letter.docx"}
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/api/export/pdf")
async def export_pdf(payload: CoverLetterPayload):
    text = payload.coverLetter.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Missing cover letter text")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=1 * inch,
        leftMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    story = []

    paragraphs = text.split("\n\n")
    for para in paragraphs:
        cleaned = para.strip().replace("\n", "<br/>")
        story.append(Paragraph(cleaned, styles["Normal"]))
        story.append(Spacer(1, 0.2 * inch))

    doc.build(story)
    buffer.seek(0)

    headers = {"Content-Disposition": "attachment; filename=cover_letter.pdf"}
    return StreamingResponse(buffer, media_type="application/pdf", headers=headers)
